#!/usr/bin/env python
"""Generate V1.5 with all 10 fixes applied."""
from docx import Document

doc = Document('d:/方法学论文/calibdeconv/manuscript/修改V1.2/CalibDeconv_submission_draft_V1_4.docx')
changes = []

def find_and_modify(search_frag, old_text, new_text):
    for p in doc.paragraphs:
        if search_frag in p.text:
            if old_text in p.text:
                modified = p.text.replace(old_text, new_text)
                p.clear(); p.add_run(modified)
                return True
    return False

def find_and_append(search_frag, append_text):
    for p in doc.paragraphs:
        if search_frag in p.text and append_text[:40] not in p.text:
            new = p.text + append_text
            p.clear(); p.add_run(new)
            return True
    return False

# 1. Methods: adaptive conformal
if find_and_append(
    'The phrase calibrated interval is used only for conformal intervals',
    ' As a sensitivity analysis, we additionally evaluated a sample-adaptive normalized-residual nonconformity score, defined as |y - yhat| / (sigma + epsilon), where sigma denotes the per-cell-type ensemble standard deviation for the given sample and epsilon = 0.001 provides numerical stability. Conformal quantiles were estimated separately for each cell type on the calibration set. Test-set prediction intervals were constructed as yhat plus or minus q times (sigma + epsilon), followed by the same boundary clipping to the unit interval.'
): changes.append('1. Methods adaptive conformal')

# 2. Adaptive tone tightened
if find_and_modify(
    'adaptive variant',
    'but the adaptive variant is available when tighter nominal tracking and sample-level interval differentiation are prioritized',
    'The adaptive normalized-residual variant served as a sensitivity analysis rather than a replacement. In this benchmark, it tracked the nominal 90% level more closely but did not improve coverage-width efficiency, producing wider intervals (0.402) while achieving lower coverage (0.901) than the fixed-radius default (0.289 width, 0.950 coverage)'
): changes.append('2. Adaptive tone tightened')

# 3. Coverage CI denominator
if find_and_modify(
    'binomial 95% confidence',
    'The empirical coverage of 0.950 at n = 500 corresponds to a binomial 95% confidence interval of [0.929, 0.967], confirming that observed coverage is statistically consistent with or above the 90% nominal target.',
    'Per-cell-type empirical coverage was computed over n = 500 held-out mixtures per lineage. At this sample size, a binomial 95% confidence interval for a point estimate of 0.950 is [0.929, 0.967], confirming that observed per-cell-type coverage is statistically above the 90% nominal target.'
): changes.append('3. Coverage CI denominator')

# 4. Supp S3 legend sync
if find_and_modify(
    'Module-level ablations',
    'Module-level ablations showing NNLS versus ensemble mean, raw ensemble interval coverage versus conformal coverage, donor-aware versus random split behavior and marker-number sensitivity.',
    'Module-level ablations showing NNLS versus ensemble mean, raw ensemble interval coverage versus fixed-radius conformal coverage, fixed-radius versus adaptive normalized-residual conformal intervals, donor-aware versus random split behavior and marker-number sensitivity.'
): changes.append('4. S3 legend synced')

# 5. S2 legend clean
if find_and_modify(
    'unavailable external tools',
    'and unavailable external tools, with method run status documented separately',
    'and a local nu-SVR implementation. Compatibility information for additional published tools is summarized in the method run-status table'
): changes.append('5. S2 legend cleaned')

# 6. Discussion opening
if find_and_modify(
    'CalibDeconv addresses a focused reliability problem',
    'CalibDeconv addresses a focused reliability problem in PBMC deconvolution.',
    'CalibDeconv treats uncertainty calibration as a first-class objective in PBMC deconvolution, rather than as an informal diagnostic attached after point estimation.'
): changes.append('6. Discussion opening')

# 7. Modularity tighten
if find_and_modify(
    'modular',
    'Importantly, the conformal calibration layer is modular: it could in principle wrap any existing point estimator, including CIBERSORTx or MuSiC, and deliver calibrated intervals on top of their predictions.',
    'Because split conformal calibration is model-agnostic, the same calibration strategy could in principle be applied to predictions from other deconvolution engines, provided that a matched calibration set with known proportions is available.'
): changes.append('7. Modularity tightened')

# 8a. Methods code
if find_and_modify(
    'will be deposited in a public repository',
    'Source code and analysis scripts will be deposited in a public repository before submission.',
    'Source code and analysis scripts are available at https://github.com/Doctor-zyx/CalibDeconv.'
): changes.append('8a. Methods code URL')

# 8b. Data availability
if find_and_modify(
    'should be deposited',
    'should be deposited with the code repository before submission',
    'are available in the code repository at https://github.com/Doctor-zyx/CalibDeconv'
): changes.append('8b. Data availability')

# 8c. Acknowledgements
if find_and_modify(
    'No specific funding',
    'No specific funding or acknowledgements are declared in this draft. This section should be updated before submission if institutional, computational or funding support needs to be acknowledged.',
    'The authors acknowledge computational resources provided by local institutional infrastructure.'
): changes.append('8c. Acknowledgements')

# 8d. Author contributions
if find_and_modify(
    'Author-specific contributions should be finalized',
    'The author(s) conceived the study, developed the method, performed the analyses, prepared the figures and wrote the manuscript. Author-specific contributions should be finalized according to the target journal requirements before submission.',
    'Y.X.Z. conceived and designed the study, developed the computational framework, performed all analyses, generated figures and wrote the manuscript.'
): changes.append('8d. Author contributions')

# 9. DC 4-type sensitivity
if find_and_append(
    'consistent with the very small number of source DCs',
    ' In a four-type sensitivity analysis excluding DC and renormalizing over T cell, B cell, NK cell and Monocyte, overall CCC increased from 0.870 to 0.974, with all per-cell-type concordances exceeding 0.97. This confirms that the weaker DC performance reflects source-cell scarcity in the external dataset rather than a calibration or signature failure in the major PBMC lineages.'
): changes.append('9. DC 4-type sensitivity')

# Save
out = 'd:/方法学论文/calibdeconv/manuscript/修改V1.2/CalibDeconv_submission_draft_V1_5.docx'
doc.save(out)
print(f'Applied {len(changes)} changes:')
for c in changes: print(f'  {c}')
print(f'\nSaved: {out}')
