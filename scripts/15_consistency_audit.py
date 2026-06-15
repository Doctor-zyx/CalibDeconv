#!/usr/bin/env python
"""Full consistency audit of V1.2 manuscript."""
from docx import Document
import re

doc = Document('d:/方法学论文/calibdeconv/manuscript/修改V1.2/CalibDeconv_submission_draft_V1_2.docx')

all_text = [p.text.strip() for p in doc.paragraphs]
fulltext = '\n'.join(all_text)

# Body = before References
body = []
for p in doc.paragraphs:
    if p.text.strip() == 'References': break
    body.append(p.text.strip())
bodytext = ' '.join(body)

print('='*70)
print('CHECK 1: Figure panels — legend vs body')
print('='*70)

# Legend panels
legends_on = False
legend_panels = {}
for t in all_text:
    if t in ('Figure legends', 'Supplementary figure legends'):
        legends_on = True; continue
    if legends_on and t.startswith('Figure '):
        fm = re.match(r'Figure (\d)\.', t)
        if fm:
            fig_num = fm.group(1)
            panels = re.findall(r'\(([A-E])\)', t)
            legend_panels[f'Fig{fig_num}'] = set(panels)

# Body panels
body_panels = {}
for m in re.finditer(r'Figure (\d)([A-E])', bodytext):
    body_panels.setdefault(f'Fig{m.group(1)}', set()).add(m.group(2))

for fig in sorted(set(list(legend_panels.keys()) + list(body_panels.keys()))):
    leg = sorted(legend_panels.get(fig, set()))
    bod = sorted(body_panels.get(fig, set()))
    status = 'OK' if set(leg) == set(bod) else 'MISMATCH'
    if status == 'MISMATCH':
        extra_leg = set(leg) - set(bod)
        extra_bod = set(bod) - set(leg)
        print(f'  {fig}: legend={leg} body={bod} -> MISMATCH')
        if extra_leg: print(f'        In legend but NOT cited: {sorted(extra_leg)}')
        if extra_bod: print(f'        Cited but NOT in legend: {sorted(extra_bod)}')
    else:
        print(f'  {fig}: {leg} OK')

print()
print('='*70)
print('CHECK 2: Supplementary figures')
print('='*70)
supp_in_legend = set()
for t in all_text:
    for m in re.finditer(r'Supplementary Figure (S\d+)', t):
        supp_in_legend.add(m.group(1))
supp_in_body = set(re.findall(r'Supplementary Fig\.\s*(S\d+)', bodytext))
print(f'  In legends: {sorted(supp_in_legend)}')
print(f'  In body:    {sorted(supp_in_body)}')
diff1 = supp_in_legend - supp_in_body
diff2 = supp_in_body - supp_in_legend
if diff1: print(f'  MISMATCH: in legend not body: {sorted(diff1)}')
if diff2: print(f'  MISMATCH: in body not legend: {sorted(diff2)}')
if not diff1 and not diff2: print('  All consistent')

print()
print('='*70)
print('CHECK 3: Terminology consistency')
print('='*70)
checks = [
    ('whole-blood vs whole blood', r'whole-blood', r'whole blood'),
    ('five-type vs 5-type', r'five-type', r'5-type'),
    ('out-of-domain vs out of domain', r'out-of-domain', r'out of domain'),
    ('pseudo-bulk vs pseudobulk', r'pseudo-bulk', r'pseudobulk'),
    ('cell-type vs cell type', r'cell-type', r'cell type'),
]
for label, pat1, pat2 in checks:
    c1 = len(re.findall(pat1, fulltext, re.I))
    c2 = len(re.findall(pat2, fulltext, re.I))
    if c1 > 0 and c2 > 0:
        print(f'  MIXED: {label} -> "{pat1}" {c1}x, "{pat2}" {c2}x')
    else:
        winner = pat1 if c1 > 0 else pat2
        print(f'  OK: uses "{winner}" consistently ({c1+c2}x)')

print()
print('='*70)
print('CHECK 4: Key numbers (abstract vs results match)')
print('='*70)
abs_text = ' '.join(all_text[:10])
nums = ['0.848', '0.844', '39%', '95%', '0.870', '0.892', '0.332', '-0.605',
        '0.627', '445', '0.081', '0.085']
for n in nums:
    in_abs = n in abs_text
    in_body = n in bodytext
    if in_abs and not in_body:
        print(f'  WARNING: "{n}" in abstract but NOT in results body')
    elif in_abs:
        print(f'  OK: "{n}" in both abstract and body')

print()
print('='*70)
print('CHECK 5: Citation order')
print('='*70)
first_app = {}
for m in re.finditer(r'\[(\d+)', bodytext):
    n = int(m.group(1))
    if n not in first_app: first_app[n] = m.start()
ordered = sorted(first_app.items(), key=lambda x: x[1])
issues = []
prev = 0
for num, pos in ordered:
    if num < prev: issues.append(f'[{num}] after [{prev}]')
    prev = max(prev, num)
if issues:
    print(f'  Out-of-order: {len(issues)} issues')
    for iss in issues[:5]: print(f'    {iss}')
else:
    print('  All citations in sequential order')

print()
print('='*70)
print('CHECK 6: Actual figure files vs legend panel count')
print('='*70)
import os
fig_dir = 'd:/方法学论文/calibdeconv/results/figures_publication_draft'
for fn in sorted(os.listdir(fig_dir)):
    if fn.endswith('.png') and 'fig' in fn.lower():
        print(f'  {fn}')
print('  (All are single-image files; panel count determined by what was drawn inside)')
print('  Figure 1 was drawn with panels A-D only (verified from image)')
print('  -> Legend (E) has NO corresponding visual panel')

print()
print('='*70)
print('SUMMARY OF ISSUES TO FIX')
print('='*70)
